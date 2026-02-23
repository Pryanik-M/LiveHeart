import io
from django.http import HttpResponse
from django.template.loader import render_to_string
from django.utils import timezone
from xhtml2pdf import pisa
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import openpyxl
from openpyxl.styles import Font, Border, Side, Alignment


# --- ВСПОМОГАТЕЛЬНЫЕ ФУНКЦИИ ---

def format_date(dt):
    """Безопасное форматирование даты"""
    if not dt:
        return "-"
    if isinstance(dt, str):
        # Если вдруг пришла строка, пробуем вернуть как есть или распарсить (но лучше парсить во view)
        return dt
    return dt.strftime('%d.%m.%Y')


def get_val(val, unit=""):
    """Если значение есть — возвращает 'Значение Ед.изм', иначе '-'"""
    if val is None or val == "":
        return "-"
    return f"{val} {unit}"


# --- ГЕНЕРАЦИЯ WORD (DOCX) ---

def generate_docx(exam):
    doc = Document()

    # Настройка стилей
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    # 1. ЗАГОЛОВОК
    header = doc.add_paragraph("ГБУЗ НО «Центральная городская больница г. Арзамас»")
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header.runs[0].bold = True

    title = doc.add_paragraph("ПРОТОКОЛ ЭХОКАРДИОГРАФИИ")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(14)

    doc.add_paragraph()  # Пустая строка

    # 2. ДАННЫЕ ПАЦИЕНТА
    p = doc.add_paragraph()
    p.add_run("Ф.И.О.: ").bold = True
    p.add_run(f"{exam.patient.full_name}\t\t")
    p.add_run("Дата: ").bold = True
    p.add_run(f"{format_date(exam.exam_datetime)}")

    p = doc.add_paragraph()
    p.add_run("Возраст: ").bold = True
    p.add_run(f"{get_val(exam.age, 'лет')}\t")
    p.add_run("Рост: ").bold = True
    p.add_run(f"{get_val(exam.height, 'см')}\t")
    p.add_run("Вес: ").bold = True
    p.add_run(f"{get_val(exam.weight, 'кг')}\t")

    p = doc.add_paragraph()
    p.add_run("ППТ: ").bold = True
    p.add_run(f"{get_val(exam.bsa, 'м²')}\t")
    p.add_run("ЧСС: ").bold = True
    p.add_run(f"{get_val(exam.hr, 'уд/мин')}")

    doc.add_paragraph("_" * 70)  # Разделитель

    # Функция для создания таблиц разделов
    def create_section_table(title, data_pairs):
        doc.add_heading(title, level=3)
        table = doc.add_table(rows=len(data_pairs), cols=2)
        table.autofit = True

        for i, (label, value) in enumerate(data_pairs):
            row = table.rows[i]
            row.cells[0].text = label
            row.cells[0].width = Inches(3.0)
            row.cells[1].text = str(value)

    # --- АОРТА ---
    if hasattr(exam, 'aorta') and exam.aorta.is_enabled:
        create_section_table("АОРТА", [
            ("Диаметр основания:", get_val(exam.aorta.diameter, "мм")),
            ("Раскрытие аортального клапана:", get_val(exam.aorta.valve_opening, "мм")),
        ])

    # --- АОРТАЛЬНЫЙ КЛАПАН ---
    if hasattr(exam, 'aorticvalve') and exam.aorticvalve.is_enabled:
        create_section_table("АОРТАЛЬНЫЙ КЛАПАН", [
            ("Пиковая скорость (Vmax):", get_val(exam.aorticvalve.psk, "м/с")),
            ("Макс. градиент давления:", get_val(exam.aorticvalve.grad_max, "мм рт.ст.")),
            ("Средний градиент:", get_val(exam.aorticvalve.grad_mean, "мм рт.ст.")),
            ("Площадь отверстия:", get_val(exam.aorticvalve.area, "см²")),
            ("Регургитация:", f"{exam.aorticvalve.regurgitation} ст."),
        ])

    # --- ЛЕВЫЙ ЖЕЛУДОЧЕК ---
    if hasattr(exam, 'leftventricle') and exam.leftventricle.is_enabled:
        lv = exam.leftventricle
        fv_val = ((lv.edv - lv.esv) / lv.edv * 100) if (lv.edv and lv.esv) else None

        create_section_table("ЛЕВЫЙ ЖЕЛУДОЧЕК", [
            ("КДР (Конечно-диаст. размер):", get_val(lv.edd, "мм")),
            ("КСР (Конечно-сист. размер):", get_val(lv.esd, "мм")),
            ("КДО (Конечно-диаст. объем):", get_val(lv.edv, "мл")),
            ("КСО (Конечно-сист. объем):", get_val(lv.esv, "мл")),
            ("МЖП (толщина в диастолу):", get_val(lv.ivsd, "мм")),
            ("ЗСЛЖ (толщина в диастолу):", get_val(lv.pw, "мм")),
            ("Фракция выброса (Simpson):", get_val(round(fv_val, 1) if fv_val else None, "%")),
        ])

    # --- ЗАКЛЮЧЕНИЕ (СЕГМЕНТЫ) ---
    doc.add_heading("ЛОКАЛЬНАЯ СОКРАТИМОСТЬ", level=3)
    p = doc.add_paragraph()
    segments = exam.segments.all().order_by('segment_number')
    norm_count = 0
    bad_segments = []

    STATES = {0: "Норма", 1: "Гипокинез", 2: "Акинез", 3: "Дискинез"}

    for s in segments:
        if s.state == 0:
            norm_count += 1
        else:
            bad_segments.append(f"Сегмент {s.segment_number}: {STATES[s.state]}")

    if norm_count == 17:
        p.add_run("Нарушения локальной сократимости не выявлены.")
    else:
        p.add_run("Выявлены зоны нарушения сократимости:\n").bold = True
        p.add_run(", ".join(bad_segments))

    # Сохранение
    f = io.BytesIO()
    doc.save(f)
    f.seek(0)
    response = HttpResponse(f.read(),
                            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="Echo_{exam.patient.full_name}.docx"'
    return response


# --- ГЕНЕРАЦИЯ EXCEL (XLSX) ---

def generate_xlsx(exam):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Протокол"

    # --- Стили ---
    thin_border = Border(left=Side(style='thin'), right=Side(style='thin'), top=Side(style='thin'),
                         bottom=Side(style='thin'))
    font_bold = Font(name='Arial', size=10, bold=True)
    font_norm = Font(name='Arial', size=10)
    font_title = Font(name='Arial', size=12, bold=True)

    # Настройка ширины колонок
    ws.column_dimensions['A'].width = 30
    ws.column_dimensions['B'].width = 15
    ws.column_dimensions['C'].width = 30
    ws.column_dimensions['D'].width = 15

    # Заголовок
    ws.merge_cells('A1:D1')
    ws['A1'] = "ГБУЗ НО «Центральная городская больница г. Арзамас»"
    ws['A1'].font = font_bold
    ws['A1'].alignment = Alignment(horizontal='center')

    ws.merge_cells('A2:D2')
    ws['A2'] = "ПРОТОКОЛ ЭХОКАРДИОГРАФИИ"
    ws['A2'].font = font_title
    ws['A2'].alignment = Alignment(horizontal='center')

    # Данные пациента (Сетка)
    row = 4

    def write_cell(r, c, val, bold=False, border=False):
        cell = ws.cell(row=r, column=c, value=val)
        cell.font = font_bold if bold else font_norm
        if border:
            cell.border = thin_border
        return cell

    write_cell(row, 1, "Ф.И.О. пациента:", bold=True)
    write_cell(row, 2, exam.patient.full_name)
    write_cell(row, 3, "Дата исследования:", bold=True)
    write_cell(row, 4, format_date(exam.exam_datetime))
    row += 1

    write_cell(row, 1, "Возраст:", bold=True)
    write_cell(row, 2, get_val(exam.age))
    write_cell(row, 3, "Рост / Вес:", bold=True)
    write_cell(row, 4, f"{get_val(exam.height)} / {get_val(exam.weight)}")
    row += 2  # Отступ

    # Функция для отрисовки разделов в рамке
    def write_section(title, data):
        nonlocal row
        # Заголовок раздела
        ws.merge_cells(f'A{row}:D{row}')
        cell = ws.cell(row=row, column=1, value=title)
        cell.font = font_bold
        cell.alignment = Alignment(horizontal='center', vertical='center')
        cell.fill = openpyxl.styles.PatternFill(start_color="EEEEEE", end_color="EEEEEE", fill_type="solid")
        for col in range(1, 5): ws.cell(row=row, column=col).border = thin_border
        row += 1

        # Данные
        for label, val in data.items():
            write_cell(row, 1, label, border=True)
            write_cell(row, 2, val, border=True)
            # Если есть вторые колонки, можно добавить логику, но пока делаем список вниз
            ws.merge_cells(f'B{row}:D{row}')  # Значение на 3 колонки для красоты
            row += 1
        row += 1

    # --- АОРТА ---
    if hasattr(exam, 'aorta') and exam.aorta.is_enabled:
        write_section("АОРТА", {
            "Диаметр основания": get_val(exam.aorta.diameter, "мм"),
            "Раскрытие створок": get_val(exam.aorta.valve_opening, "мм")
        })

    # --- ЛЕВЫЙ ЖЕЛУДОЧЕК ---
    if hasattr(exam, 'leftventricle') and exam.leftventricle.is_enabled:
        lv = exam.leftventricle
        fv_val = ((lv.edv - lv.esv) / lv.edv * 100) if (lv.edv and lv.esv) else None
        write_section("ЛЕВЫЙ ЖЕЛУДОЧЕК", {
            "КДР": get_val(lv.edd, "мм"),
            "КСР": get_val(lv.esd, "мм"),
            "КДО": get_val(lv.edv, "мл"),
            "КСО": get_val(lv.esv, "мл"),
            "Толщина МЖП": get_val(lv.ivsd, "мм"),
            "Толщина ЗСЛЖ": get_val(lv.pw, "мм"),
            "Фракция выброса (Simpson)": get_val(round(fv_val, 1) if fv_val else None, "%")
        })

    # Сохранение
    f = io.BytesIO()
    wb.save(f)
    f.seek(0)
    response = HttpResponse(f.read(), content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = f'attachment; filename="Echo_{exam.patient.full_name}.xlsx"'
    return response


# --- ГЕНЕРАЦИЯ PDF ---

def generate_pdf(exam):
    # Используем шаблон pdf_report.html
    html_string = render_to_string('patients/pdf_report.html', {'exam': exam})
    result = io.BytesIO()
    # Поддержка кириллицы требует шрифтов, но xhtml2pdf имеет встроенные ограничения.
    # Для базовой работы убедитесь, что в HTML есть <meta charset="utf-8">
    pdf = pisa.pisaDocument(io.BytesIO(html_string.encode("UTF-8")), result)

    if not pdf.err:
        response = HttpResponse(result.getvalue(), content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="Echo_{exam.patient.full_name}.pdf"'
        return response
    return HttpResponse("Ошибка PDF")